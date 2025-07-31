import re
from pathlib import Path
from typing import Dict, Any, List
from pydantic import Field
from docx import Document
from langchain.tools import BaseTool
from function import iter_block_items, _copy_run_format
from docx.table import Table
from docx.text.paragraph import Paragraph

"""
Word文档处理器
@author QianTianhao
@date 2025-07-16
"""

class WordDocumentProcessor:


    def __init__(self, output_dir: str = "./chapters"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.chapter_names = []

    def extract_chapters(self, doc_path: str) -> List[str]:
        doc = Document(doc_path)
        chapters = {}
        current = None
        content = []

        for block in iter_block_items(doc):
            text = block.text.strip() if hasattr(block, "text") else ""
            if isinstance(block, Paragraph) and self._is_chapter_heading(block):
                if current and content:
                    chapters[current] = content
                current = self._clean_chapter_name(block.text.strip())
                content = [block]
            elif current:
                content.append(block)
        if current and content:
            chapters[current] = content

        # 保存章节
        for name, items in chapters.items():
            self._save_chapter(name, items)
        self.chapter_names = list(chapters.keys())
        return self.chapter_names

    def _is_chapter_heading(self, paragraph) -> bool:
        """判断段落是否为章节标题 - 强制只匹配汉字数字章节"""
        text = paragraph.text.strip()
        if not text:
            return False

        # 首先检查是否为目录项（目录项通常较短且可能包含页码）
        if self._is_toc_item(text):
            return False

        # 只检查汉字数字章节格式，严格匹配
        chinese_chapter_patterns = [
            r'^第[一二三四五六七八九十百千万]+章\s+.+',  # 第一章 标题
            r'^第[一二三四五六七八九十百千万]+章$',  # 单独的第一章
        ]

        for pattern in chinese_chapter_patterns:
            if re.match(pattern, text):
                # 进一步验证：章节标题通常不会太长，且确保不包含阿拉伯数字
                if len(text) <= 100 and not re.search(r'\d', text):
                    return True

        return False

    def _is_toc_item(self, text: str) -> bool:
        """判断是否为目录项"""
        # 目录项的特征
        toc_patterns = [
            r'\.{3,}',  # 多个点号
            r'\d+$',  # 以数字结尾（页码）
            r'\.\.\.',  # 省略号
            r'^\d+\s*$',  # 单独的数字
        ]

        for pattern in toc_patterns:
            if re.search(pattern, text):
                return True

        # 目录项通常包含特定词汇且格式特殊
        if '目录' in text or 'contents' in text.lower():
            return True

        # 如果文本很短且只包含数字和基本字符，可能是目录
        if len(text) < 10 and re.match(r'^[\d\s\.\-]+$', text):
            return True

        return False

    def _clean_chapter_name(self, text: str) -> str:
        """清理章节名称，用作文件名"""
        # 移除或替换Windows文件名中的非法字符
        illegal_chars = {
            '<': '＜', '>': '＞', ':': '：', '"': '＂',
            '/': '／', '\\': '＼', '|': '｜', '?': '？', '*': '＊'
        }

        cleaned = text
        for old_char, new_char in illegal_chars.items():
            cleaned = cleaned.replace(old_char, new_char)

            # 压缩多个空格
            cleaned = ' '.join(cleaned.split())

            # **去掉末尾的常见标点**：
            cleaned = cleaned.rstrip('；;：:、,，.。')

            # 限制文件名长度
            if len(cleaned) > 50:
                cleaned = cleaned[:50].rstrip()

            return cleaned

    def _save_chapter(self, chapter_name: str, items: List[Any]) -> str:
        new_doc = Document()
        for block in items:
            if isinstance(block, Paragraph):
                new_para = new_doc.add_paragraph()
                for run in block.runs:
                    new_run = new_para.add_run(run.text)
                    try:
                        _copy_run_format(run, new_run)
                    except Exception:
                        # 即便失败也不要中断
                        pass
            elif isinstance(block, Table):
                tbl = new_doc.add_table(rows=0, cols=len(block.columns))
                # 复制每行
                for row in block.rows:
                    new_row = tbl.add_row().cells
                    for idx, cell in enumerate(row.cells):
                        new_row[idx].text = cell.text
        path = self.output_dir / f"{chapter_name}.docx"
        new_doc.save(path)
        return str(path)

"""
增强的章节阅读工具，支持读取段落和表格内容
@author QianTianhao
@date 2025-07-16
"""
class ChapterReaderTool(BaseTool):
    name: str = "chapter_reader"
    description: str = "读取指定章节的内容，包括段落和表格。输入应为章节名称的字符串。"
    chapters_dir: Path = Field(default_factory=lambda: Path("./chapters"))

    class Config:
        arbitrary_types_allowed = True

    def __init__(self, chapters_dir: str = "./chapters", **kwargs):
        super().__init__(chapters_dir=Path(chapters_dir), **kwargs)

    def _run(self, chapter_name: str) -> str:
        """读取章节内容，包括段落和表格，使用鲁棒的表格提取方法"""
        file_path = self.chapters_dir / f"{chapter_name}.docx"

        if not file_path.exists():
            return f"章节文件不存在: {chapter_name}"

        try:
            doc = Document(file_path)
            content = []

            # 提取段落（排除表格内容）
            table_paragraphs = set()
            try:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                table_paragraphs.add(paragraph._element)
            except:
                pass

            # 读取段落
            for paragraph in doc.paragraphs:
                if paragraph._element not in table_paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        content.append(text)

            # 使用鲁棒方法提取表格
            tables_data = self._extract_tables_robust(doc)

            for table_data in tables_data:
                table_content = self._format_table_content_robust(table_data)
                if table_content:
                    content.append(f"\n[表格 {table_data['index'] + 1}]")
                    content.append(table_content)
                    content.append("[表格结束]\n")

            result = "\n".join(content)

            # 调试输出
            print(f"=== 读取章节: {chapter_name} ===")
            print(f"内容长度: {len(result)} 字符")
            print(f"段落数: {len(doc.paragraphs)}")
            print(f"表格数: {len(tables_data)}")
            print(f"前500字符: {result[:500]}...")
            print("=" * 50)

            return result

        except Exception as e:
            error_msg = f"读取章节失败: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            return error_msg

    def _extract_tables_robust(self, doc) -> list:
        """高成功率的表格提取方法"""
        tables_data = []

        # 方法1：使用python-docx的标准方法
        try:
            for i, table in enumerate(doc.tables):
                table_data = self._extract_single_table_standard(table, i)
                if table_data:
                    tables_data.append(table_data)
        except Exception as e:
            print(f"标准方法提取表格失败: {e}")

        # 方法2：如果标准方法失败，使用XML解析
        if not tables_data:
            try:
                tables_data = self._extract_tables_from_xml(doc)
            except Exception as e:
                print(f"XML方法提取表格失败: {e}")

        # 方法3：使用文本模式识别
        if not tables_data:
            try:
                tables_data = self._extract_tables_from_text(doc)
            except Exception as e:
                print(f"文本方法提取表格失败: {e}")

        return tables_data

    def _extract_single_table_standard(self, table, table_index) -> dict:
        """使用标准方法提取单个表格"""
        try:
            if not table or not table.rows:
                return None

            rows_data = []
            max_cols = 0

            for row_idx, row in enumerate(table.rows):
                row_data = []
                try:
                    for cell in row.cells:
                        cell_text = self._extract_cell_text_robust(cell)
                        row_data.append(cell_text)

                    if row_data:
                        rows_data.append(row_data)
                        max_cols = max(max_cols, len(row_data))

                except Exception as e:
                    print(f"读取表格 {table_index} 第 {row_idx} 行失败: {e}")
                    # 尝试从XML直接读取这一行
                    try:
                        row_data = self._extract_row_from_xml(row)
                        if row_data:
                            rows_data.append(row_data)
                    except:
                        pass

            # 标准化行长度
            for row in rows_data:
                while len(row) < max_cols:
                    row.append("")

            return {
                'index': table_index,
                'rows': rows_data,
                'row_count': len(rows_data),
                'col_count': max_cols
            }

        except Exception as e:
            print(f"标准方法提取表格 {table_index} 失败: {e}")
            return None

    def _extract_cell_text_robust(self, cell) -> str:
        """鲁棒的单元格文本提取"""
        try:
            # 方法1：标准提取
            texts = []
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()
                if text:
                    texts.append(text)

            if texts:
                return " ".join(texts)

            # 方法2：从XML提取
            try:
                cell_xml = cell._element
                text_content = self._extract_text_from_xml(cell_xml)
                if text_content:
                    return text_content
            except:
                pass

            # 方法3：尝试从runs提取
            all_text = ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        all_text += run.text

            return all_text.strip()

        except Exception as e:
            print(f"提取单元格文本失败: {e}")
            return ""

    def _extract_row_from_xml(self, row) -> list:
        """从XML提取行数据"""
        try:
            row_data = []
            for cell in row.cells:
                cell_text = self._extract_cell_text_robust(cell)
                row_data.append(cell_text)
            return row_data
        except:
            return []

    def _extract_tables_from_xml(self, doc) -> list:
        """从XML直接提取表格"""
        try:
            tables_data = []
            document_xml = doc.element

            # 查找所有table元素
            tables = document_xml.xpath('.//w:tbl', namespaces={
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

            for table_idx, table_elem in enumerate(tables):
                rows_data = []

                # 查找所有行
                rows = table_elem.xpath('.//w:tr', namespaces={
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

                for row_elem in rows:
                    row_data = []

                    # 查找所有单元格
                    cells = row_elem.xpath('.//w:tc', namespaces={
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

                    for cell_elem in cells:
                        cell_text = self._extract_text_from_xml(cell_elem)
                        row_data.append(cell_text)

                    if row_data:
                        rows_data.append(row_data)

                if rows_data:
                    tables_data.append({
                        'index': table_idx,
                        'rows': rows_data,
                        'row_count': len(rows_data),
                        'col_count': max(len(row) for row in rows_data) if rows_data else 0
                    })

            return tables_data

        except Exception as e:
            print(f"XML提取表格失败: {e}")
            return []

    def _extract_text_from_xml(self, element) -> str:
        """从XML元素提取文本"""
        try:
            # 查找所有文本节点
            text_nodes = element.xpath('.//w:t',
                                       namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            texts = [node.text for node in text_nodes if node.text]
            return " ".join(texts).strip()
        except:
            return ""

    def _extract_tables_from_text(self, doc) -> list:
        """从文档文本中提取表格结构"""
        try:
            tables_data = []
            full_text = ""

            # 获取完整文档文本
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"

            lines = full_text.split('\n')
            current_table = []

            for line in lines:
                line = line.strip()
                if not line:
                    if current_table:
                        # 结束当前表格
                        if len(current_table) > 1:  # 至少2行才算表格
                            tables_data.append({
                                'index': len(tables_data),
                                'rows': current_table,
                                'row_count': len(current_table),
                                'col_count': max(len(row) for row in current_table) if current_table else 0
                            })
                        current_table = []
                    continue

                # 检查是否包含表格分隔符
                if '\t' in line or '|' in line or '，' in line:
                    # 尝试不同的分隔符
                    separators = ['\t', '|', '，', '  ']
                    for sep in separators:
                        if sep in line:
                            row_data = [cell.strip() for cell in line.split(sep)]
                            if len(row_data) > 1:  # 至少2列
                                current_table.append(row_data)
                                break

            # 处理最后一个表格
            if current_table and len(current_table) > 1:
                tables_data.append({
                    'index': len(tables_data),
                    'rows': current_table,
                    'row_count': len(current_table),
                    'col_count': max(len(row) for row in current_table) if current_table else 0
                })

            return tables_data

        except Exception as e:
            print(f"文本提取表格失败: {e}")
            return []

    def _format_table_content_robust(self, table_data) -> str:
        """格式化表格内容（鲁棒版本）"""
        try:
            if not table_data or not table_data.get('rows'):
                return ""

            rows = table_data['rows']
            table_lines = []

            # 检测表头
            has_header = len(rows) > 1 and self._detect_header_from_data(rows)

            for i, row in enumerate(rows):
                if not row or not any(str(cell).strip() for cell in row):
                    continue

                row_text = " | ".join(str(cell).strip() for cell in row)
                table_lines.append(row_text)

                # 在表头后添加分隔线
                if i == 0 and has_header:
                    table_lines.append("-" * len(row_text))

            return "\n".join(table_lines)

        except Exception as e:
            print(f"格式化表格失败: {e}")
            return f"[表格 {table_data.get('index', '?')} - 格式化失败]"

    def _detect_header_from_data(self, rows) -> bool:
        """从数据推断是否有表头"""
        if len(rows) < 2:
            return False

        try:
            first_row = rows[0]
            second_row = rows[1]

            # 检查第一行是否主要是文字，第二行是否有数字
            first_row_numeric = sum(
                1 for cell in first_row if str(cell).strip().replace('.', '').replace('-', '').isdigit())
            second_row_numeric = sum(
                1 for cell in second_row if str(cell).strip().replace('.', '').replace('-', '').isdigit())

            # 如果第一行数字少，第二行数字多，可能是表头
            if len(first_row) > 0 and len(second_row) > 0:
                first_ratio = first_row_numeric / len(first_row)
                second_ratio = second_row_numeric / len(second_row)
                return first_ratio < 0.3 and second_ratio > 0.3

            return False

        except:
            return False

    def extract_cell_text(self, cell) -> str:
        """提取单元格中的所有文本"""
        texts = []
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if text:
                texts.append(text)
        return "\n".join(texts)
